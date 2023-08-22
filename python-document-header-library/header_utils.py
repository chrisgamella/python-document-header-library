from docx import Document

class HeaderGenerator:
    def __init__(self):
        self.doc = Document()

    def add_paragraph(self, content):
        self.doc.add_paragraph(content)

    def save(self, file_path):
        self.doc.save(file_path)
